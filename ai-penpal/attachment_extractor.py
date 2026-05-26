"""
attachment_extractor.py
-----------------------
Utilities for extracting plain text from supported email attachments.

Image math extraction prefers Ollama + LLaVA. When vision extraction is not
available, the extractor falls back to pytesseract OCR for general image text.
Tesseract requires the binary to be installed separately on the host machine:
  macOS:  brew install tesseract
  Ubuntu: sudo apt install tesseract-ocr
"""

from io import BytesIO
from pathlib import Path
import re
import tempfile

import logging

from config import (
    ENABLE_IMAGE_OCR,
    MAX_EXTRACTED_CHARS_PER_ATTACHMENT,
    SUPPORTED_ATTACHMENT_EXTENSIONS,
)
from math_validator import (
    detect_math_expression,
    is_expression_suspicious,
    normalize_math_expression,
    safe_evaluate_expression,
)
from vision_math import extract_math_from_image

logger = logging.getLogger(__name__)

try:
    from pypdf import PdfReader
except ImportError:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from PIL import Image, UnidentifiedImageError
except ImportError:
    Image = None

    class UnidentifiedImageError(Exception):
        pass

try:
    import pytesseract
    from pytesseract import TesseractNotFoundError
except ImportError:
    pytesseract = None

    class TesseractNotFoundError(Exception):
        pass


CONTENT_TYPE_EXTENSION_MAP = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "image/png": ".png",
    "image/jpeg": ".jpg",
}
AMBIGUOUS_MATH_CHARS = set("Il|OoSs")


def extract_text_from_attachment(filename: str, content_type: str, file_bytes: bytes) -> dict:
    """Extract text from a supported attachment and never raise."""
    resolved_filename = filename or "attachment"
    resolved_content_type = content_type or "application/octet-stream"
    extension = _detect_extension(resolved_filename, resolved_content_type)
    result = {
        "filename": resolved_filename,
        "content_type": resolved_content_type,
        "supported": False,
        "text": "",
        "error": None,
        "math_expression": None,
        "math_expression_uncertain": False,
        "math_expression_note": None,
        "computed_result": None,
        "math_confirmation_expression": None,
        "normalized_math_expression": None,
        "math_source": None,
        "math_confidence": None,
    }

    if extension not in SUPPORTED_ATTACHMENT_EXTENSIONS:
        return result

    result["supported"] = True

    if not file_bytes:
        result["error"] = "Attachment was empty."
        return result

    try:
        if extension == ".pdf":
            text = _extract_pdf_text(file_bytes)
        elif extension == ".docx":
            text = _extract_docx_text(file_bytes)
        elif extension in {".png", ".jpg", ".jpeg"}:
            text, math_metadata = _extract_image_content(file_bytes, extension)
            result.update(math_metadata)
        else:
            return result

        result["text"] = _truncate_text(text)
        return result
    except Exception as exc:
        logger.warning(
            "[ATTACHMENT] Extraction failed for %s (%s): %s",
            resolved_filename,
            resolved_content_type,
            exc,
        )
        result["error"] = str(exc)
        return result


def _detect_extension(filename: str, content_type: str) -> str:
    extension = Path(filename).suffix.lower()
    if extension:
        return extension
    return CONTENT_TYPE_EXTENSION_MAP.get(content_type.lower(), "")


def _truncate_text(text: str) -> str:
    normalized = (text or "").strip()
    if len(normalized) <= MAX_EXTRACTED_CHARS_PER_ATTACHMENT:
        return normalized
    return normalized[:MAX_EXTRACTED_CHARS_PER_ATTACHMENT].rstrip()


def _extract_pdf_text(file_bytes: bytes) -> str:
    if PdfReader is None:
        raise RuntimeError("PDF extraction library not installed.")

    reader = PdfReader(BytesIO(file_bytes))
    chunks = []
    total_chars = 0

    for index, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:
            logger.warning("[ATTACHMENT] Skipped page %s due to error: %s", index, exc)
            continue
        if not page_text:
            continue
        remaining = MAX_EXTRACTED_CHARS_PER_ATTACHMENT - total_chars
        if remaining <= 0:
            break
        chunk = page_text[:remaining]
        chunks.append(chunk)
        total_chars += len(chunk)

    return "\n".join(chunks)


def _extract_docx_text(file_bytes: bytes) -> str:
    if Document is None:
        raise RuntimeError("python-docx is not installed.")

    document = Document(BytesIO(file_bytes))
    chunks = []
    total_chars = 0

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()
        if not paragraph_text:
            continue
        remaining = MAX_EXTRACTED_CHARS_PER_ATTACHMENT - total_chars
        if remaining <= 0:
            break
        chunk = paragraph_text[:remaining]
        chunks.append(chunk)
        total_chars += len(chunk)

    if total_chars < MAX_EXTRACTED_CHARS_PER_ATTACHMENT:
        for table in document.tables:
            for row in table.rows:
                cell_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if not cell_text:
                    continue
                remaining = MAX_EXTRACTED_CHARS_PER_ATTACHMENT - total_chars
                if remaining <= 0:
                    break
                chunk = cell_text[:remaining]
                chunks.append(chunk)
                total_chars += len(chunk)
            if total_chars >= MAX_EXTRACTED_CHARS_PER_ATTACHMENT:
                break

    return "\n".join(chunks)


def _extract_image_text(file_bytes: bytes) -> str:
    if not ENABLE_IMAGE_OCR:
        raise RuntimeError("Image OCR is disabled in configuration.")
    if Image is None or pytesseract is None:
        raise RuntimeError("Image OCR dependencies are not installed.")

    try:
        with Image.open(BytesIO(file_bytes)) as image:
            prepared = image.convert("RGB")
            return pytesseract.image_to_string(prepared)
    except TesseractNotFoundError as exc:
        raise RuntimeError(
            "Tesseract OCR is not installed on this machine. "
            "Install it with 'brew install tesseract' on macOS or "
            "'sudo apt install tesseract-ocr' on Ubuntu."
        ) from exc
    except UnidentifiedImageError as exc:
        raise RuntimeError("Attachment is not a readable image.") from exc


def _extract_image_content(file_bytes: bytes, extension: str) -> tuple[str, dict]:
    vision_result = _extract_image_math_with_vision(file_bytes, extension)
    if vision_result is not None:
        logger.info("[ATTACHMENT] Vision math extraction succeeded for image attachment")
        return vision_result

    raw_text = _extract_image_text(file_bytes)
    return _format_image_ocr_text(raw_text)


def _extract_image_math_with_vision(file_bytes: bytes, extension: str) -> tuple[str, dict] | None:
    tmp_path = None

    try:
        with tempfile.NamedTemporaryFile(suffix=extension, delete=False) as tmp_file:
            tmp_file.write(file_bytes)
            tmp_path = tmp_file.name

        vision_result = extract_math_from_image(tmp_path)
        expression = (vision_result.get("expression") or "").strip()
        if not expression or not detect_math_expression(expression):
            return None

        validation = safe_evaluate_expression(expression)
        normalized_expression = validation["expression"] or normalize_math_expression(expression)
        if validation["error"] or validation["result"] is None or not normalized_expression:
            logger.info("[ATTACHMENT] Vision math extraction returned an unusable expression")
            return None

        metadata = {
            "math_expression": expression,
            "math_expression_uncertain": False,
            "math_expression_note": None,
            "computed_result": validation["result"],
            "math_confirmation_expression": None,
            "normalized_math_expression": normalized_expression,
            "math_source": "vision",
            "math_confidence": vision_result.get("confidence", "low"),
        }
        text = (
            f"Interpreted expression from the image:\n"
            f"{expression}\n\n"
            f"Normalized:\n"
            f"{normalized_expression}\n\n"
            f"Computed result:\n"
            f"{validation['result']}"
        )
        return text, metadata
    finally:
        if tmp_path:
            Path(tmp_path).unlink(missing_ok=True)


def _format_image_ocr_text(raw_text: str) -> tuple[str, dict]:
    metadata = {
        "math_expression": None,
        "math_expression_uncertain": False,
        "math_expression_note": None,
        "computed_result": None,
        "math_confirmation_expression": None,
        "normalized_math_expression": None,
        "math_source": "ocr",
        "math_confidence": None,
    }

    if not detect_math_expression(raw_text):
        return raw_text, metadata

    validation = safe_evaluate_expression(raw_text)
    normalized_expression = validation["expression"] or normalize_math_expression(raw_text)
    confirmation_expression = _suggest_confirmation_expression(normalized_expression)
    uncertain = _math_ocr_is_uncertain(
        raw_text,
        normalized_expression,
        validation["error"],
        confirmation_expression,
    )
    computed_result = validation["result"] if not uncertain else None

    lines = [
        "OCR detected math expression:",
        normalized_expression or raw_text.strip(),
        "Computed result:",
        computed_result if computed_result is not None else "Unavailable",
    ]

    note = None
    if uncertain or validation["error"]:
        note = "OCR extraction may be inaccurate. Please verify the expression from the image."
        lines.append(f"Note: {note}")

    metadata.update(
        {
            "math_expression": normalized_expression or raw_text.strip(),
            "math_expression_uncertain": uncertain or validation["error"] is not None,
            "math_expression_note": note,
            "computed_result": computed_result,
            "math_confirmation_expression": confirmation_expression,
            "normalized_math_expression": normalized_expression or raw_text.strip(),
            "math_source": "ocr",
            "math_confidence": None,
        }
    )

    return "\n".join(lines).strip(), metadata


def _math_ocr_is_uncertain(
    raw_text: str,
    normalized_expression: str,
    error: str | None,
    confirmation_expression: str | None,
) -> bool:
    if error:
        return True
    if any(char in AMBIGUOUS_MATH_CHARS for char in raw_text):
        return True
    if normalized_expression.count("(") != normalized_expression.count(")"):
        return True
    if not normalized_expression:
        return True
    if is_expression_suspicious(normalized_expression):
        return True
    if confirmation_expression and confirmation_expression != _display_math_expression(normalized_expression):
        return True
    if not _math_confidence_established(raw_text, normalized_expression):
        return True
    return False


def _math_confidence_established(raw_text: str, normalized_expression: str) -> bool:
    raw_math_chars = re.sub(r"[^0-9+\-*/=().xX×÷−–—]", "", raw_text)
    normalized_math_chars = re.sub(r"[^0-9+\-*/().]", "", normalized_expression)

    if not raw_math_chars or not normalized_math_chars:
        return False

    similarity = len(set(raw_math_chars) & set(normalized_math_chars)) / max(len(set(normalized_math_chars)), 1)
    if similarity < 0.6:
        return False

    operator_count = sum(1 for char in normalized_expression if char in "+-*/")
    return operator_count >= 1


def _suggest_confirmation_expression(normalized_expression: str) -> str:
    corrected_expression = normalized_expression

    if re.match(r"^\d+\+\d+\*\([^)]+\)$", corrected_expression):
        corrected_expression = corrected_expression.replace("+", "/", 1)

    corrected_expression = re.sub(
        r"\((\d)(\d)([+\-])(\d)\)",
        lambda match: f"({match.group(1)}{match.group(3)}{match.group(4)})",
        corrected_expression,
    )

    return _display_math_expression(corrected_expression)


def _display_math_expression(expression: str) -> str:
    display = expression
    display = re.sub(r"\*(?=\()", "", display)
    display = display.replace("/", " ÷ ")
    display = display.replace("*", " × ")
    display = re.sub(r"\s+", " ", display).strip()
    return display
